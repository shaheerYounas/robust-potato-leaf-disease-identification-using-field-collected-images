"""
Convert Markdown report to DOCX using BIP journal template formatting.

Formatting requirements from Final Report Template.docx:
- Times New Roman font throughout
- A4 paper (21cm x 29.7cm)
- Top/Left margins 2.5 cm, Bottom/Right margins 2.0 cm
- 1.15 line spacing
- Title: 12pt Bold, Sentence case
- Section headings: 12pt Bold
- Body text: 12pt
- Tables/figures captions: 10pt
- Abstract, keywords: 11pt
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


def setup_document():
    """Create a blank document with BIP template formatting."""
    doc = Document()

    # Page setup: A4, margins
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Set line spacing to 1.15
    pf = style.paragraph_format
    pf.line_spacing = 1.15

    return doc


def add_title(doc, text):
    """Add title: 12pt Bold, centered."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    p.paragraph_format.line_spacing = 1.15


def add_author_block(doc, lines):
    """Add author/affiliation block centered."""
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        if "Author" in line:
            run.bold = True
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(0)


def add_section_heading(doc, text):
    """Add section heading: 12pt Bold, all caps per template."""
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(12)


def add_subsection_heading(doc, text):
    """Add subsection heading: 12pt Bold, normal case."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(6)


def add_body_paragraph(doc, text):
    """Add body text: 12pt, justified, 1.15 spacing."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Handle bold and italic inline formatting
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def add_keywords(doc, text):
    """Add keywords line: 11pt."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    p.paragraph_format.line_spacing = 1.15


def add_table(doc, header_row, data_rows, caption=None):
    """Add a formatted table with optional caption above."""
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        run.bold = True
        p.paragraph_format.space_after = Pt(2)

    num_cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = "Table Grid"

    # Header
    for i, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text.strip())
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

    # Data rows
    for r_idx, row_data in enumerate(data_rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text.strip())
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"

    # Add spacing after table
    doc.add_paragraph()


def parse_markdown_table(lines):
    """Parse a Markdown table into header and data rows."""
    header = None
    data = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if all(c.replace("-", "").replace(":", "").strip() == "" for c in cells):
            continue  # separator row
        if header is None:
            header = cells
        else:
            data.append(cells)
    return header, data


def convert_md_to_docx(md_path, docx_path):
    """Convert a Markdown report file to DOCX with BIP template formatting."""
    md_text = Path(md_path).read_text(encoding="utf-8")
    lines = md_text.split("\n")

    doc = setup_document()

    i = 0
    in_table = False
    table_lines = []
    table_caption = None

    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            if in_table and table_lines:
                header, data = parse_markdown_table(table_lines)
                if header:
                    add_table(doc, header, data, table_caption)
                table_lines = []
                in_table = False
                table_caption = None
            i += 1
            continue

        # Title (# )
        if line.startswith("# ") and not line.startswith("## "):
            add_title(doc, line[2:].strip())
            i += 1
            continue

        # Author block with ** markers
        if line.startswith("**Author"):
            author_lines = []
            while i < len(lines) and lines[i].strip():
                clean = lines[i].strip().replace("**", "")
                author_lines.append(clean)
                i += 1
            add_author_block(doc, author_lines)
            continue

        # Affiliation lines (¹, ², e-mail:)
        if line.strip().startswith("¹") or line.strip().startswith("²") or line.strip().startswith("e-mail:"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip())
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # Section heading (## )
        if line.startswith("## "):
            heading_text = line[3:].strip()
            add_section_heading(doc, heading_text)
            i += 1
            continue

        # Subsection heading (### )
        if line.startswith("### "):
            heading_text = line[4:].strip()
            add_subsection_heading(doc, heading_text)
            i += 1
            continue

        # Table caption (**Table N.**)
        if line.strip().startswith("**Table"):
            table_caption = line.strip().replace("**", "")
            i += 1
            continue

        # Table row
        if line.strip().startswith("|"):
            in_table = True
            table_lines.append(line)
            i += 1
            continue

        # List items starting with -
        if line.strip().startswith("- "):
            p = doc.add_paragraph()
            p.style = doc.styles["List Bullet"]
            text = line.strip()[2:]
            # Handle bold
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
            p.paragraph_format.line_spacing = 1.15
            i += 1
            continue

        # Regular paragraph
        add_body_paragraph(doc, line.strip())
        i += 1

    # Flush remaining table
    if in_table and table_lines:
        header, data = parse_markdown_table(table_lines)
        if header:
            add_table(doc, header, data, table_caption)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    base = Path(__file__).resolve().parent.parent
    reports_dir = base / "submission_ready" / "final_package" / "reports"

    convert_md_to_docx(
        reports_dir / "internship_report_submission.md",
        reports_dir / "internship_report_submission.docx",
    )
    convert_md_to_docx(
        reports_dir / "journal_paper_submission.md",
        reports_dir / "journal_paper_submission.docx",
    )
    print("Done — both DOCX files generated.")
